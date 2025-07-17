import openai
import pandas as pd
from dotenv import load_dotenv
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from ttkbootstrap import Style
from ttkbootstrap.constants import *
from ttkbootstrap import ttk
from threading import Thread
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Pt

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY") # API Key da OpenAI

# Função principal para gerar o plano de teste
def gerar_casos():
    tarefa = entry_tarefa.get("1.0", tk.END)
    contexto = entry_contexto.get()
    diretorio = entry_diretorio.get()

    if not tarefa or not diretorio:
        messagebox.showerror("Erro", "Preencha todos os campos obrigatórios.")
        return
    
    progress_bar.start()

    def gerador():
        try:
            prompt = f"""
            Você é um Tester, Especialista em Quality Assurance e precisa criar um plano de teste para a tarefa: '{tarefa}'. '{contexto}'.

            Forneça:

            1. Resumo da Tarefa
            Resumo breve da tarefa descrita

            2. Pre-condições
            Liste os dados ou pré-condições necessárias para a execução dos testes focando na experiência do usuário.

            3. Tecnicas de Teste
            Liste as técnicas de teste que serão utilizadas

            4. Casos de Teste
            Gere os casos de testes, utilizando tecnicas como: Valor limite, Estado, Particionamento de Equivalência entre outras. Cenários positivos e negativos.
            Para criação dos casos e retorne no formato CSV com as colunas:  
            Caso de Teste;Descrição;Passos;Resultado Esperado;Status;Observação do teste.

            Obs: "Caso de teste" deve ser retorntado no formato CT-00,
            "Passos" deve retonrar em BDD,
            "Status" e "Observação do teste" deve ser preenchido com -
            Apenas Isso, não precisa retornar mais nada.
            """

            response = openai.chat.completions.create(
                model="gpt-4.1", # Versão do modelo da OpenAI // Ajustar para mais recente se necessário
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1 # Temperatura // baixa para respostas mais objetivas e alta para mais criativas
            )

            conteudo = response.choices[0].message.content

            partes = conteudo.split("Caso de Teste;")
            cabecalho_extra = partes[0].strip()
            csv_content = "Caso de Teste;" + partes[1].strip()

            from io import StringIO
            df = pd.read_csv(StringIO(csv_content), delimiter=';')

            timestamp = datetime.now().strftime("%Y%d%m_%H%M%S")
            doc_filename = os.path.join(diretorio, f"plano_de_teste_{timestamp}.docx")
            document = Document()

            # Alterando a orientação da página para paisagem
            section = document.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_height, section.page_width = section.page_width, section.page_height  # Ajuste de largura e altura

            # Cabeçalho com logo
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = 1

            logo_path = os.path.join(os.path.dirname(__file__), "logo.png")  # Logo
            if os.path.exists(logo_path):
                run = header_para.add_run()
                run.add_picture(logo_path, width=Inches(1.5))

            # Adicionando título em negrito e centralizado
            titulo = document.add_paragraph('PLANO DE TESTES')
            titulo.alignment = 1  # Centralizar
            run = titulo.runs[0]
            run.bold = True  # Colocar em negrito

            table_info = document.add_table(rows=4, cols=2)
            table_info.style = 'Table Grid'
            table_info.cell(0, 0).text = "Data do plano:"
            table_info.cell(0, 1).text = datetime.now().strftime("%d/%m/%Y")
            table_info.cell(1, 0).text = "Data do teste:"
            table_info.cell(1, 1).text = ""
            table_info.cell(2, 0).text = "Testador:"
            table_info.cell(2, 1).text = ""
            table_info.cell(3, 0).text = "Versão:"
            table_info.cell(3, 1).text = ""

            document.add_paragraph('')
            document.add_heading("Tarefa", level=1)
            document.add_paragraph(tarefa)  # Tarefa com formatação preservada

            if contexto.strip():
                document.add_heading("Contexto", level=1)
                document.add_paragraph(contexto)

            document.add_paragraph('')
            document.add_heading("Resumo", level=1)
            document.add_paragraph(cabecalho_extra)

            document.add_paragraph('')
            document.add_heading("Casos de Teste", level=1)

            table = document.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = col

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

            document.save(doc_filename)

        except Exception as e:
            messagebox.showerror("Erro", str(e))
        finally:
            progress_bar.stop()

    Thread(target=gerador).start()

def selecionar_diretorio():
    pasta = filedialog.askdirectory(initialdir=os.path.expanduser("~/Downloads"))
    if pasta:
        entry_diretorio.delete(0, tk.END)
        entry_diretorio.insert(0, pasta)

# Interface Gráfica
app = tk.Tk()
app.title("Gerador de Plano de Teste por IA")
app.geometry("700x500")
style = Style("flatly")

frame = ttk.Frame(app, padding=20)
frame.pack(fill=tk.BOTH, expand=True)

ttk.Label(frame, text="Descrição da Tarefa *").pack(anchor=tk.W)
entry_tarefa = tk.Text(frame, height=6, wrap=tk.WORD)
entry_tarefa.pack(fill=tk.X, pady=5)

ttk.Label(frame, text="Contexto (opcional)").pack(anchor=tk.W)
entry_contexto = ttk.Entry(frame)
entry_contexto.pack(fill=tk.X, pady=5)

ttk.Label(frame, text="Diretório para salvar *").pack(anchor=tk.W)
dir_frame = ttk.Frame(frame)
dir_frame.pack(fill=tk.X, pady=5)

entry_diretorio = ttk.Entry(dir_frame)
entry_diretorio.pack(side=tk.LEFT, fill=tk.X, expand=True)
entry_diretorio.insert(0, os.path.expanduser("~/Downloads"))
ttk.Button(dir_frame, text="Selecionar", command=selecionar_diretorio).pack(side=tk.RIGHT, padx=5)

progress_bar = ttk.Progressbar(frame, mode="indeterminate")
progress_bar.pack(fill=tk.X, pady=10)

ttk.Button(frame, text="Gerar Plano de Teste", command=gerar_casos, bootstyle=SUCCESS).pack(pady=10)

app.mainloop()
